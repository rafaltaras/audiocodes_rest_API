import datetime, os, base64, configparser, requests, json, re, string
from docx import Document
from docx.shared import Pt
from docx.shared import Inches

class Config:
    def __init__(self):
        pass

    def ensure_path(self, folderpath):   
        if os.path.exists(folderpath) is not True:
            try:
                os.mkdir(folderpath)
            except OSError:
                print(f"Creation file filed {folderpath}")
            else:
                print(f"File created {folderpath}")
        return folderpath

    def make_filepath(self, ip, folderpath):
        dt = str(datetime.datetime.now())
        filepath = folderpath + ip + "_" + dt[:10] 
        return filepath
    
    def get_url_ini_file(self, ip):
        url = f'http://{ip}/api/v1/files/ini'
        return url

    def credential(self, username, password, url):
        cred = username + ':' + password
        cred_encoded = base64.b64encode(cred.encode()).decode()
        headers = {'Authorization': 'Basic ' + cred_encoded} 
        return requests.get(url, headers=headers)

    def download_config(self, ip, username, password):
        folderpath = self.ensure_path("./backups/")        
        filepath = self.make_filepath(ip, folderpath)
        url = self.get_url_ini_file(ip)
        response = self.credential(username, password, url)
        if response.status_code != 200:
            print("Backup filed: ", response.status_code)
            return None
        filepath = filepath + ".ini"
        with open(filepath, "w") as f:
            f.write(response.text)
        print(filepath)
        return filepath

    def read_ini(self, path): 
        config = configparser.ConfigParser()
        config.sections()
        config.read(path)
        return config
    
    def filter_sections(self, config):
        sections = []
        for section in config.sections():
            if not section.startswith(" \\"):
                sections.append(section)
        return sections

    def read_parameter_details(self, path, section):
        config = self.read_ini(path)
        return config[section]

    def create_docx_file(self, ip, path, sections):
        document = Document()
        folderpath = self.ensure_path("./backups/")        
        filepath = self.make_filepath(ip, folderpath)
        paragraph = document.add_paragraph(f"Configuration of AudioCodes {ip}")
        paragraph.style = document.styles['Heading 1']
        for section in sections:
            document.add_heading(section)
            parameters_details = self.read_parameter_details(path, section)
            for key, value in parameters_details.items():
                table = document.add_table(rows=1, cols=2)
                cells = table.rows[0].cells
                cells[0].text = key
                cells[1].text = value
        document.save(f'{filepath}'+'.docx')
        return None

    def create_ini_dict(self, parameter_details):
        param_details = {}
        for k,v in parameter_details.items():
            v = v.split(',')
            param_details.setdefault(k,v) 
        return param_details

config  = Config()